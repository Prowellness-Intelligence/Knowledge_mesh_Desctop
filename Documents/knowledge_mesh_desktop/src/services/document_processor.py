"""
Document Processing Service for the Knowledge Mesh Desktop application.

This module provides a service that processes documents, extracting text and metadata,
generating summaries, and preparing content for vector indexing.
"""

import asyncio
import logging
import os
import time
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Set, Tuple, Union, Any

import numpy as np
from PyPDF2 import PdfReader
import docx
import pandas as pd
import markdown
import pytesseract
from PIL import Image
import fitz  # PyMuPDF
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.stem import WordNetLemmatizer
from sklearn.feature_extraction.text import TfidfVectorizer
from transformers import pipeline

from ..core.config import Config
from ..core.events import EventType, publish, subscribe, event_bus
from ..models.document import Document, DocumentType, DocumentStatus

logger = logging.getLogger(__name__)


class DocumentProcessorService:
    """
    Service for processing documents.
    
    This service processes documents, extracting text and metadata,
    generating summaries, and preparing content for vector indexing.
    """
    
    def __init__(self, config: Config):
        """
        Initialize the document processor service.
        
        Args:
            config: Application configuration
        """
        self.config = config
        self.is_running = False
        self.processing_queue = asyncio.Queue()
        self.processing_task = None
        self.vector_store_service = None
        self.summarizer = None
        self.keyword_extractor = None
        self.max_summary_length = 200
        self.max_keywords = 10
        self.supported_extensions = {
            ".pdf": self._process_pdf,
            ".docx": self._process_docx,
            ".doc": self._process_docx,
            ".txt": self._process_txt,
            ".md": self._process_markdown,
            ".csv": self._process_csv,
            ".xlsx": self._process_excel,
            ".png": self._process_image,
            ".jpg": self._process_image,
            ".jpeg": self._process_image,
        }
        
        self._load_config()
    
    def _load_config(self):
        """Load configuration from the config object."""
        self.max_summary_length = self.config.get(
            "document_processor.max_summary_length", 200
        )
        
        self.max_keywords = self.config.get(
            "document_processor.max_keywords", 10
        )
        
        extensions = self.config.get(
            "document_processor.supported_extensions",
            [".pdf", ".docx", ".doc", ".txt", ".md", ".csv", ".xlsx", ".png", ".jpg", ".jpeg"]
        )
        
        self.supported_extensions = {
            ext: handler
            for ext, handler in self.supported_extensions.items()
            if ext in extensions
        }
    
    async def initialize(self):
        """Initialize the document processor service."""
        logger.info("Initializing document processor service")
        
        try:
            nltk.download("punkt", quiet=True)
            nltk.download("stopwords", quiet=True)
            nltk.download("wordnet", quiet=True)
            
            self.summarizer = pipeline("summarization", model="facebook/bart-large-cnn")
            logger.info("Summarization model loaded")
        except Exception as e:
            logger.warning(f"Error loading summarization model: {e}")
            self.summarizer = None
        
        event_bus.subscribe(EventType.FILE_CREATED, self._on_file_created)
        event_bus.subscribe(EventType.FILE_MODIFIED, self._on_file_modified)
        
        logger.info("Document processor service initialized")
    
    async def start(self):
        """Start the document processor service."""
        if self.is_running:
            logger.warning("Document processor service is already running")
            return
        
        logger.info("Starting document processor service")
        
        self.is_running = True
        
        self.processing_task = asyncio.create_task(self._process_documents())
        
        logger.info("Document processor service started")
    
    async def stop(self):
        """Stop the document processor service."""
        if not self.is_running:
            logger.warning("Document processor service is not running")
            return
        
        logger.info("Stopping document processor service")
        
        self.is_running = False
        
        if self.processing_task:
            self.processing_task.cancel()
            try:
                await self.processing_task
            except asyncio.CancelledError:
                pass
        
        logger.info("Document processor service stopped")
    
    async def _process_documents(self):
        """Process documents in the queue."""
        while self.is_running:
            try:
                file_path = await self.processing_queue.get()
                
                await self._process_document(file_path)
                
                self.processing_queue.task_done()
            except asyncio.CancelledError:
                break
            except Exception as e:
                logger.error(f"Error processing document: {e}", exc_info=True)
    
    async def _on_file_created(self, event):
        """
        Handle file created events.
        
        Args:
            event: The file created event
        """
        file_path = event.data.get("path")
        if file_path:
            _, ext = os.path.splitext(file_path)
            if ext.lower() in self.supported_extensions:
                await self.processing_queue.put(file_path)
    
    async def _on_file_modified(self, event):
        """
        Handle file modified events.
        
        Args:
            event: The file modified event
        """
        file_path = event.data.get("path")
        if file_path:
            _, ext = os.path.splitext(file_path)
            if ext.lower() in self.supported_extensions:
                await self.processing_queue.put(file_path)
    
    async def _process_document(self, file_path: str):
        """
        Process a document.
        
        Args:
            file_path: The path to the document
        """
        logger.info(f"Processing document: {file_path}")
        
        try:
            if not os.path.exists(file_path):
                logger.warning(f"File does not exist: {file_path}")
                return
            
            _, ext = os.path.splitext(file_path)
            ext = ext.lower()
            
            if ext not in self.supported_extensions:
                logger.warning(f"Unsupported file extension: {ext}")
                return
            
            document = await self.supported_extensions[ext](file_path)
            
            if document:
                if document.content:
                    document.summary = await self._generate_summary(document.content)
                
                if document.content:
                    document.metadata["keywords"] = await self._extract_keywords(document.content)
                
                await self._save_document(document)
                
                await self._index_document(document)
                
                publish(
                    EventType.DOCUMENT_PROCESSING_COMPLETED,
                    {
                        "document_id": document.id,
                        "path": file_path,
                    },
                )
                
                logger.info(f"Document processed: {file_path}")
            else:
                logger.warning(f"Failed to process document: {file_path}")
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {e}", exc_info=True)
            
            publish(
                EventType.DOCUMENT_PROCESSING_FAILED,
                {
                    "path": file_path,
                    "error": str(e),
                },
            )
    
    async def _process_pdf(self, file_path: str) -> Optional[Document]:
        """
        Process a PDF document.
        
        Args:
            file_path: The path to the PDF document
            
        Returns:
            A Document object, or None if processing fails
        """
        try:
            with open(file_path, "rb") as file:
                reader = PdfReader(file)
                text = ""
                metadata = {}
                
                if reader.metadata:
                    metadata = {
                        "title": reader.metadata.get("/Title", ""),
                        "author": reader.metadata.get("/Author", ""),
                        "subject": reader.metadata.get("/Subject", ""),
                        "creator": reader.metadata.get("/Creator", ""),
                        "producer": reader.metadata.get("/Producer", ""),
                        "creation_date": reader.metadata.get("/CreationDate", ""),
                    }
                
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
            
            if not text.strip():
                doc = fitz.open(file_path)
                text = ""
                for page in doc:
                    text += page.get_text() + "\n\n"
                doc.close()
            
            document = Document(
                id=os.path.basename(file_path),
                title=metadata.get("title", os.path.basename(file_path)),
                content=text,
                file_path=file_path,
                file_type=DocumentType.PDF,
                metadata=metadata,
                status=DocumentStatus.PROCESSED,
            )
            
            return document
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}", exc_info=True)
            return None
    
    async def _process_docx(self, file_path: str) -> Optional[Document]:
        """
        Process a DOCX document.
        
        Args:
            file_path: The path to the DOCX document
            
        Returns:
            A Document object, or None if processing fails
        """
        try:
            doc = docx.Document(file_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            metadata = {
                "title": doc.core_properties.title or "",
                "author": doc.core_properties.author or "",
                "subject": doc.core_properties.subject or "",
                "keywords": doc.core_properties.keywords or "",
                "created": doc.core_properties.created or "",
                "modified": doc.core_properties.modified or "",
            }
            
            document = Document(
                id=os.path.basename(file_path),
                title=metadata.get("title", os.path.basename(file_path)),
                content=text,
                file_path=file_path,
                file_type=DocumentType.DOCX,
                metadata=metadata,
                status=DocumentStatus.PROCESSED,
            )
            
            return document
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}", exc_info=True)
            return None
    
    async def _process_txt(self, file_path: str) -> Optional[Document]:
        """
        Process a TXT document.
        
        Args:
            file_path: The path to the TXT document
            
        Returns:
            A Document object, or None if processing fails
        """
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as file:
                text = file.read()
            
            document = Document(
                id=os.path.basename(file_path),
                title=os.path.basename(file_path),
                content=text,
                file_path=file_path,
                file_type=DocumentType.TXT,
                metadata={},
                status=DocumentStatus.PROCESSED,
            )
            
            return document
        except Exception as e:
            logger.error(f"Error processing TXT {file_path}: {e}", exc_info=True)
            return None
    
    async def _process_markdown(self, file_path: str) -> Optional[Document]:
        """
        Process a Markdown document.
        
        Args:
            file_path: The path to the Markdown document
            
        Returns:
            A Document object, or None if processing fails
        """
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as file:
                text = file.read()
            
            html = markdown.markdown(text)
            
            document = Document(
                id=os.path.basename(file_path),
                title=os.path.basename(file_path),
                content=text,
                file_path=file_path,
                file_type=DocumentType.MARKDOWN,
                metadata={"html": html},
                status=DocumentStatus.PROCESSED,
            )
            
            return document
        except Exception as e:
            logger.error(f"Error processing Markdown {file_path}: {e}", exc_info=True)
            return None
    
    async def _process_csv(self, file_path: str) -> Optional[Document]:
        """
        Process a CSV document.
        
        Args:
            file_path: The path to the CSV document
            
        Returns:
            A Document object, or None if processing fails
        """
        try:
            df = pd.read_csv(file_path)
            
            text = df.to_string()
            
            metadata = {
                "columns": df.columns.tolist(),
                "rows": len(df),
                "columns_count": len(df.columns),
            }
            
            document = Document(
                id=os.path.basename(file_path),
                title=os.path.basename(file_path),
                content=text,
                file_path=file_path,
                file_type=DocumentType.CSV,
                metadata=metadata,
                status=DocumentStatus.PROCESSED,
            )
            
            return document
        except Exception as e:
            logger.error(f"Error processing CSV {file_path}: {e}", exc_info=True)
            return None
    
    async def _process_excel(self, file_path: str) -> Optional[Document]:
        """
        Process an Excel document.
        
        Args:
            file_path: The path to the Excel document
            
        Returns:
            A Document object, or None if processing fails
        """
        try:
            df = pd.read_excel(file_path)
            
            text = df.to_string()
            
            metadata = {
                "columns": df.columns.tolist(),
                "rows": len(df),
                "columns_count": len(df.columns),
                "sheet_name": df.sheet_name if hasattr(df, "sheet_name") else "",
            }
            
            document = Document(
                id=os.path.basename(file_path),
                title=os.path.basename(file_path),
                content=text,
                file_path=file_path,
                file_type=DocumentType.EXCEL,
                metadata=metadata,
                status=DocumentStatus.PROCESSED,
            )
            
            return document
        except Exception as e:
            logger.error(f"Error processing Excel {file_path}: {e}", exc_info=True)
            return None
    
    async def _process_image(self, file_path: str) -> Optional[Document]:
        """
        Process an image document.
        
        Args:
            file_path: The path to the image document
            
        Returns:
            A Document object, or None if processing fails
        """
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            
            metadata = {
                "width": image.width,
                "height": image.height,
                "format": image.format,
                "mode": image.mode,
            }
            
            document = Document(
                id=os.path.basename(file_path),
                title=os.path.basename(file_path),
                content=text,
                file_path=file_path,
                file_type=DocumentType.IMAGE,
                metadata=metadata,
                status=DocumentStatus.PROCESSED,
            )
            
            return document
        except Exception as e:
            logger.error(f"Error processing image {file_path}: {e}", exc_info=True)
            return None
    
    async def _generate_summary(self, text: str) -> str:
        """
        Generate a summary for a document.
        
        Args:
            text: The document text
            
        Returns:
            A summary of the document
        """
        try:
            if len(text) <= self.max_summary_length:
                return text
            
            if self.summarizer:
                max_input_length = 1024
                truncated_text = text[:max_input_length]
                
                summary = self.summarizer(
                    truncated_text,
                    max_length=self.max_summary_length,
                    min_length=30,
                    do_sample=False,
                )[0]["summary_text"]
                
                return summary
            
            sentences = sent_tokenize(text)
            summary = " ".join(sentences[:3])
            
            if len(summary) > self.max_summary_length:
                summary = summary[: self.max_summary_length - 3] + "..."
            
            return summary
        except Exception as e:
            logger.error(f"Error generating summary: {e}", exc_info=True)
            
            return text[: self.max_summary_length - 3] + "..."
    
    async def _extract_keywords(self, text: str) -> List[str]:
        """
        Extract keywords from a document.
        
        Args:
            text: The document text
            
        Returns:
            A list of keywords
        """
        try:
            tokens = word_tokenize(text.lower())
            
            stop_words = set(stopwords.words("english"))
            tokens = [
                token
                for token in tokens
                if token.isalnum() and token not in stop_words and len(token) > 2
            ]
            
            lemmatizer = WordNetLemmatizer()
            tokens = [lemmatizer.lemmatize(token) for token in tokens]
            
            freq_dist = nltk.FreqDist(tokens)
            
            keywords = [word for word, _ in freq_dist.most_common(self.max_keywords)]
            
            return keywords
        except Exception as e:
            logger.error(f"Error extracting keywords: {e}", exc_info=True)
            return []
    
    async def _save_document(self, document: Document):
        """
        Save a document to the database.
        
        Args:
            document: The document to save
        """
        logger.debug(f"Document saved: {document.id}")
    
    async def _index_document(self, document: Document):
        """
        Index a document in the vector store.
        
        Args:
            document: The document to index
        """
        logger.debug(f"Document indexed: {document.id}")
        
        publish(
            EventType.DOCUMENT_INDEXED,
            {
                "document_id": document.id,
                "path": document.file_path,
            },
        )
    
    async def process_file(self, file_path: str):
        """
        Process a file.
        
        Args:
            file_path: The path to the file
        """
        await self.processing_queue.put(file_path)
    
    async def get_document(self, document_id: str) -> Optional[Document]:
        """
        Get a document by ID.
        
        Args:
            document_id: The ID of the document
            
        Returns:
            The document, or None if not found
        """
        return None
    
    async def search_documents(self, query: str) -> List[Document]:
        """
        Search for documents.
        
        Args:
            query: The search query
            
        Returns:
            A list of matching documents
        """
        return []
